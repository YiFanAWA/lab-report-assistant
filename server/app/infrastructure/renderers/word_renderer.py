"""Word 文档渲染器。

从已确认大纲生成 .docx 文件。
使用 python-docx 库，模板驱动。

设计要点：
- 从同一份已确认大纲生成，不直接从模型临时上下文生成
- 章节按顺序渲染：封面、实验目的、实验背景、数据描述、分析方案、实验结果、结论与讨论
- 执行产物（CSV/PNG）按 source_ids 关联到对应章节
- 输出文件路径由调用方指定，渲染器只负责生成文件
- SPEC 0010：支持项目级 Word 模板（Jinja2 风格占位符）
  - 封面变量：{{project_name}} {{project_topic}} {{generated_date}}
  - 章节循环：{{#sections}}...{{/sections}}
  - 循环内变量：{{section_title}} {{section_content}} {{section_source_type}} {{section_source_ids}}
  - 模板解析失败时降级到默认渲染
"""

import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from app.core.errors import AppError
from app.modules.outlines.layout_planner import (
    LayoutKind,
    SectionLayoutPlan,
    plan_section_layouts,
)
from app.modules.outlines.document_planner import (
    ThesisChapterPlan,
    plan_thesis_document,
)

logger = logging.getLogger(__name__)

# Jinja2 风格占位符正则
_VAR_PATTERN = re.compile(r"\{\{\s*(\w+)\s*\}\}")
_SECTION_START = "{{#sections}}"
_SECTION_END = "{{/sections}}"

# SPEC 0031：论文式文档视觉 token。所有默认 Word 输出均从这里取值，
# 避免在章节、图表和表格渲染方法中散落格式常量。
PAGE_MARGIN_LEFT_CM = 2.4
PAGE_MARGIN_RIGHT_CM = 2.4
PAGE_MARGIN_TOP_CM = 2.2
PAGE_MARGIN_BOTTOM_CM = 2.2
BODY_FONT_CN = "宋体"
BODY_FONT_EN = "Times New Roman"
ACCENT_HEX = "1F4E79"
TEXT_HEX = "222222"
MUTED_HEX = "6B7280"
LIGHT_ACCENT_HEX = "EAF1F8"
TABLE_HEADER_HEX = "1F4E79"
TABLE_ALT_HEX = "F5F7FA"
MAX_TABLE_ROWS = 10
MAX_TABLE_COLS = 6


class WordRenderer:
    """Word 文档渲染器。

    从已确认大纲的 sections 列表生成 .docx 文件。
    支持项目级模板（SPEC 0010）。
    """

    def render(
        self,
        project_name: str,
        project_topic: str,
        outline_sections: list[dict],
        execution_artifacts: list[dict],
        output_path: str,
        config: dict[str, Any] | None = None,
    ) -> str:
        """渲染 Word 文档（默认渲染，无模板）。

        参数：
        - project_name: 项目名称（用于封面）
        - project_topic: 项目课题（用于封面）
        - outline_sections: 已确认大纲的 sections 列表（dict 形式）
        - execution_artifacts: 执行产物列表（含 file_path/name/artifact_type）
        - output_path: 输出文件绝对路径

        返回：生成的文件路径

        异常：渲染失败抛出 AppError(code="WORD_RENDER_FAILED")。
        """
        try:
            config = config or {}
            self._formal_paper = config.get("document_profile") == "formal_academic"
            self._reader_first_paper = self._formal_paper and bool(
                config.get("reader_first", True)
            )
            self._include_audit_appendix = bool(
                config.get("include_audit_appendix", False)
            )
            doc = Document()
            self._figure_index = 0
            self._table_index = 0
            self._chapter_figure_counts: dict[int, int] = {}
            self._chapter_table_counts: dict[int, int] = {}
            self._citation_map: dict[str, int] = {}
            self._bookmark_id = 0
            document_plan = plan_thesis_document(
                project_topic,
                outline_sections,
                execution_artifacts,
                formal=self._formal_paper,
                reference_catalog=config.get("reference_catalog"),
                abstract_override=config.get("abstract"),
                abstract_en_override=config.get("abstract_en"),
                abstract_sections_override=config.get("abstract_sections"),
                abstract_sections_en_override=config.get("abstract_sections_en"),
                formal_title=config.get("formal_title", ""),
                formal_subtitle=config.get("formal_subtitle", ""),
                formal_metadata=config.get("formal_metadata"),
            )
            self._publication_profile = document_plan.publication_profile
            self._reader_first_paper = self._formal_paper and bool(
                config.get(
                    "reader_first",
                    getattr(self._publication_profile, "reader_first", True),
                )
            )
            self._include_audit_appendix = bool(
                config.get(
                    "include_audit_appendix",
                    getattr(self._publication_profile, "include_audit_appendix", False),
                )
            )
            self._enforce_formal_sufficiency(document_plan)
            self._citation_map = dict(document_plan.citation_map)
            self._configure_document(doc, project_name)

            # 封面
            self._render_cover(doc, project_name, project_topic, document_plan)
            if self._formal_paper:
                self._add_formal_section(doc, page_format="lowerRoman")
                self._render_front_matter(doc, document_plan)
                self._render_formal_figure_table_catalog(doc, execution_artifacts, document_plan)
                self._add_formal_section(doc, page_format="decimal")
            else:
                self._render_front_matter(doc, document_plan)

            # 正文：一级章节由文档规划器统一分组，二级标题保留原始大纲语义。
            for chapter in document_plan.chapters:
                self._render_chapter(doc, chapter, execution_artifacts)

            self._render_references(doc, document_plan.references)
            self._render_appendix(
                doc, execution_artifacts,
            )

            # 确保输出目录存在
            output = Path(output_path)
            output.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output))
            return str(output)
        except AppError:
            raise
        except Exception as exc:
            raise AppError(
                code="WORD_RENDER_FAILED",
                message=f"Word 文档生成失败：{exc}",
            ) from exc

    def render_with_template(
        self,
        template_path: str,
        project_name: str,
        project_topic: str,
        outline_sections: list[dict],
        execution_artifacts: list[dict],
        output_path: str,
    ) -> str:
        """使用项目级模板渲染 Word 文档（SPEC 0010）。

        流程：
        1. 打开模板 .docx
        2. 识别 {{#sections}}...{{/sections}} 循环块
        3. 循环块外：替换封面变量
        4. 循环块内：按每个 section 重复段落，替换章节变量
        5. 执行产物按 source_ids 嵌入到对应章节
        6. 保存到输出路径

        降级策略：
        - 模板文件不存在 → WORD_TEMPLATE_FILE_MISSING
        - 模板无法打开 → WORD_TEMPLATE_PARSE_FAILED
        - 循环标记不匹配 → WORD_TEMPLATE_SECTION_BLOCK_INVALID
        """
        template = Path(template_path)
        if not template.exists():
            raise AppError(
                code="WORD_TEMPLATE_FILE_MISSING",
                message=f"模板文件不存在：{template_path}",
            )

        try:
            doc = Document(str(template))
        except Exception as exc:
            raise AppError(
                code="WORD_TEMPLATE_PARSE_FAILED",
                message=f"模板文件无法打开：{exc}",
            ) from exc

        try:
            self._figure_index = 0
            self._table_index = 0
            # 准备封面变量
            cover_vars = {
                "project_name": project_name or "",
                "project_topic": project_topic or "",
                "generated_date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
                "student_name": "",
                "course_name": "",
            }

            # 识别章节循环块（在段落级别）
            start_idx, end_idx = self._find_section_block(doc.paragraphs)

            if start_idx is not None and end_idx is None:
                # 有开始标记但无结束标记
                raise AppError(
                    code="WORD_TEMPLATE_SECTION_BLOCK_INVALID",
                    message="模板含 {{#sections}} 但缺少 {{/sections}}",
                )

            if start_idx is not None and end_idx is not None:
                # 替换循环块
                self._render_template_sections(
                    doc, start_idx, end_idx, outline_sections,
                    execution_artifacts, cover_vars,
                )
                # 循环块外（封面等）的变量也需要替换
                self._replace_cover_vars(doc, cover_vars)
            else:
                # 无循环块：只替换封面变量
                self._replace_cover_vars(doc, cover_vars)
                # 追加默认章节渲染（保持兼容）
                for section in outline_sections:
                    self._render_section(doc, section, execution_artifacts)

            # 确保输出目录存在
            output = Path(output_path)
            output.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output))
            return str(output)
        except AppError:
            raise
        except Exception as exc:
            raise AppError(
                code="WORD_RENDER_FAILED",
                message=f"模板渲染失败：{exc}",
            ) from exc

    # --- 模板辅助方法 ---

    def _find_section_block(self, paragraphs) -> tuple[int | None, int | None]:
        """在段落中查找 {{#sections}} 和 {{/sections}} 标记的位置。

        返回 (start_idx, end_idx)，未找到返回 (None, None)。
        """
        start_idx = None
        end_idx = None
        for i, p in enumerate(paragraphs):
            text = p.text
            if _SECTION_START in text and start_idx is None:
                start_idx = i
            if _SECTION_END in text and start_idx is not None:
                end_idx = i
                break
        return start_idx, end_idx

    def _replace_cover_vars(self, doc: Document, cover_vars: dict) -> None:
        """替换文档中所有段落的封面变量。"""
        for paragraph in doc.paragraphs:
            self._replace_vars_in_paragraph(paragraph, cover_vars)
        # 表格中的段落也替换
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_vars_in_paragraph(paragraph, cover_vars)

    def _replace_vars_in_paragraph(self, paragraph, vars_dict: dict) -> None:
        """替换单个段落中的变量。"""
        if not paragraph.runs:
            # 无 runs 的段落直接替换文本
            if paragraph.text:
                new_text = self._replace_vars(paragraph.text, vars_dict)
                paragraph.text = new_text
            return
        # 有 runs 的段落：保留第一个 run 的格式，替换文本
        full_text = paragraph.text
        if not full_text:
            return
        new_text = self._replace_vars(full_text, vars_dict)
        if new_text != full_text:
            # 清空所有 runs，只在第一个 run 写入新文本
            first_run = paragraph.runs[0]
            first_run.text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

    def _replace_vars(self, text: str, vars_dict: dict) -> str:
        """替换文本中的 {{var}} 占位符。"""
        def replacer(match):
            var_name = match.group(1)
            return str(vars_dict.get(var_name, match.group(0)))
        return _VAR_PATTERN.sub(replacer, text)

    def _render_template_sections(
        self,
        doc: Document,
        start_idx: int,
        end_idx: int,
        outline_sections: list[dict],
        execution_artifacts: list[dict],
        cover_vars: dict,
    ) -> None:
        """渲染章节循环块。

        策略：使用文本重建方式。
        1. 收集循环块外的段落文本（保留封面等）
        2. 收集循环块内的段落文本作为模板
        3. 删除所有段落
        4. 按"循环块前 + 每个section的循环块内容 + 循环块后"重建段落
        5. 执行产物追加到末尾
        """
        paragraphs = list(doc.paragraphs)

        # 收集循环块前的段落文本（不含 start 标记）
        before_lines = [paragraphs[i].text for i in range(0, start_idx)]

        # 收集循环块内的段落文本（不含 start/end 标记）
        template_lines = [paragraphs[i].text for i in range(start_idx + 1, end_idx)]

        # 收集循环块后的段落文本（不含 end 标记）
        after_lines = [paragraphs[i].text for i in range(end_idx + 1, len(paragraphs))]

        # 删除所有现有段落
        body = doc.element.body
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)

        # 重建：循环块前的段落
        for line in before_lines:
            doc.add_paragraph(line)

        # 为每个 section 生成循环块内的段落
        for section in outline_sections:
            section_vars = {
                "section_title": section.get("title", ""),
                "section_content": section.get("content", ""),
                "section_source_type": section.get("source_type", ""),
                "section_source_ids": ", ".join(
                    str(sid) for sid in section.get("source_ids", []) or []
                ),
            }
            merged_vars = {**cover_vars, **section_vars}

            for line in template_lines:
                new_text = self._replace_vars(line, merged_vars)
                doc.add_paragraph(new_text)

        # 重建：循环块后的段落
        for line in after_lines:
            doc.add_paragraph(line)

        # 执行产物嵌入（按 source_ids 关联到章节）
        for section in outline_sections:
            source_type = section.get("source_type", "")
            source_ids = section.get("source_ids", []) or []
            if source_type == "EXECUTION":
                relevant = [
                    art for art in execution_artifacts
                    if not source_ids or art.get("execution_run_id") in source_ids
                ]
                if relevant:
                    doc.add_heading("执行产物", level=2)
                    for art in relevant:
                        self._render_single_artifact(doc, art)

    def _render_single_artifact(
        self, doc: Document, art: dict, *, chapter_number: int | None = None,
        keep_note_with_next: bool = False,
    ) -> None:
        """渲染单个执行产物，使用论文式图题/表题和可用宽度适配。"""
        name = art.get("name", "")
        file_path = art.get("file_path", "")
        art_type = art.get("artifact_type", "")

        if art_type == "CHART_PNG" and file_path:
            path = Path(file_path)
            if path.exists():
                try:
                    self._figure_index += 1
                    figure_label = f"图 {self._figure_index}"
                    image_paragraph = doc.add_paragraph()
                    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    image_paragraph.paragraph_format.space_before = Pt(8)
                    image_paragraph.paragraph_format.space_after = Pt(2)
                    image_paragraph.paragraph_format.keep_with_next = True
                    run = image_paragraph.add_run()
                    run.add_picture(
                        str(path),
                        width=Inches(
                            self._image_width_inches(
                                path, doc, art.get("figure_layout_profile"),
                            )
                        ),
                    )
                    caption = doc.add_paragraph(style="Caption")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.keep_with_next = True
                    if getattr(self, "_formal_paper", False):
                        self._add_bookmarked_sequence(
                            caption,
                            f"fig_{self._figure_index}",
                            "Figure",
                            figure_label,
                        )
                        caption.add_run(
                            f"  {art.get('figure_caption') or self._display_artifact_name(name)}"
                        ).bold = True

                    else:
                        caption.add_run(
                            f"{figure_label}  {art.get('figure_caption') or self._display_artifact_name(name)}"
                        ).bold = True
                    if getattr(self, "_reader_first_paper", False):
                        note_text = str(art.get("figure_note") or "").strip()
                        if note_text:
                            note = doc.add_paragraph(style="Figure Note")
                            note.paragraph_format.keep_with_next = keep_note_with_next
                            note.add_run(note_text)
                        # 图题和图注已经承担图像识别与解释责任；不在每张图后重复
                        # 插入“相关结果见图……”句，避免正文被机械交叉引用打断。
                    else:
                        source = doc.add_paragraph(style="Artifact Source")
                        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 图题、来源和论证说明必须作为一个连续单元分页，避免图留在上一页、
                        # 论证说明孤零地落到下一页，破坏论文阅读节奏。
                        source.paragraph_format.keep_with_next = True
                        source.add_run(art.get("figure_note") or f"来源：执行产物 {name}")
                        asset_ids = art.get("scientific_asset_ids", []) or []
                        attributions = art.get("scientific_asset_attributions", []) or []
                        if asset_ids:
                            source.add_run(
                                f"；科研资产：{', '.join(str(value) for value in asset_ids)}"
                            )
                        for attribution in attributions:
                            source.add_run(f"；{attribution}")
                        if art.get("scientific_asset_image_sha256"):
                            source.add_run(
                                f"；图像 SHA-256：{art['scientific_asset_image_sha256']}"
                            )
                        if art.get("scientific_asset_render_metadata"):
                            source.add_run(
                                f"；渲染追溯：{art['scientific_asset_render_metadata']}"
                            )
                        argument = art.get("figure_argument") or {}
                        if argument:
                            argument_note = doc.add_paragraph(style="Artifact Source")
                            argument_note.paragraph_format.keep_together = True
                            argument_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            argument_note.add_run("论证：").bold = True
                            argument_note.add_run(
                                f"主张：{argument.get('claim', '')}；"
                                f"方法：{argument.get('method', '')}；"
                                f"结果：{argument.get('result', '')}；"
                                f"边界：{argument.get('boundary', '')}。"
                            )
                            if art.get("figure_body_reference"):
                                argument_note.add_run(
                                    f"（正文引用：{art['figure_body_reference']}）"
                                )
                except Exception:
                    self._add_artifact_error(doc, f"图片无法嵌入：{name}")
            else:
                self._add_artifact_error(doc, f"图片文件不存在：{name}")
        elif art_type == "TABLE_CSV" and file_path:
            path = Path(file_path)
            if path.exists():
                try:
                    self._render_csv_table(
                        doc,
                        path,
                        chapter_number=chapter_number,
                        caption_text=art.get("table_caption"),
                        keep_note_with_next=keep_note_with_next,
                    )
                except Exception:
                    self._add_artifact_error(doc, f"表格无法嵌入：{name}")
            else:
                if getattr(self, "_formal_paper", False):
                    caption = doc.add_paragraph(style="Caption")
                    caption.paragraph_format.keep_with_next = True
                    self._add_bookmarked_sequence(
                        caption, f"tbl_{self._table_index + 1}", "Table",
                        f"表 {self._table_index + 1}",
                    )
                    caption.add_run("  " + (art.get("table_caption") or self._display_artifact_name(name)))
                self._add_artifact_error(doc, f"表格文件不存在：{name}")

    def _configure_document(self, doc: Document, project_name: str) -> None:
        """配置页面、样式和页眉页脚；正式论文版式来自 ManuscriptPlan。"""
        profile = getattr(self, "_publication_profile", None)
        is_formal = getattr(self, "_formal_paper", False)
        margin_left = float(getattr(profile, "page_margin_left_cm", PAGE_MARGIN_LEFT_CM)) if is_formal else PAGE_MARGIN_LEFT_CM
        margin_right = float(getattr(profile, "page_margin_right_cm", PAGE_MARGIN_RIGHT_CM)) if is_formal else PAGE_MARGIN_RIGHT_CM
        margin_top = float(getattr(profile, "page_margin_top_cm", PAGE_MARGIN_TOP_CM)) if is_formal else PAGE_MARGIN_TOP_CM
        margin_bottom = float(getattr(profile, "page_margin_bottom_cm", PAGE_MARGIN_BOTTOM_CM)) if is_formal else PAGE_MARGIN_BOTTOM_CM
        self._body_font_cn = getattr(profile, "body_font_cjk", BODY_FONT_CN) if is_formal else BODY_FONT_CN
        self._body_font_en = getattr(profile, "body_font_latin", BODY_FONT_EN) if is_formal else BODY_FONT_EN
        self._paper_text_color = getattr(profile, "text_color_hex", TEXT_HEX) if is_formal else TEXT_HEX
        self._paper_muted_color = getattr(profile, "muted_color_hex", MUTED_HEX) if is_formal else MUTED_HEX
        self._paper_accent_color = (
            self._paper_text_color
            if is_formal and getattr(profile, "formal_monochrome", True)
            else ACCENT_HEX
        )
        body_size = float(getattr(profile, "body_size_pt", 10.5)) if is_formal else 12
        self._body_size_pt = body_size
        self._title_size_pt = float(getattr(profile, "title_size_pt", 22.0)) if is_formal else 26
        self._subtitle_size_pt = float(getattr(profile, "subtitle_size_pt", 13.0)) if is_formal else 14
        self._heading1_size_pt = float(getattr(profile, "heading1_size_pt", 16.0)) if is_formal else 18
        self._heading2_size_pt = float(getattr(profile, "heading2_size_pt", 12.0)) if is_formal else 14.5
        self._heading3_size_pt = float(getattr(profile, "heading3_size_pt", 10.5)) if is_formal else 12
        self._caption_size_pt = float(getattr(profile, "caption_size_pt", 9.0)) if is_formal else 9
        self._table_size_pt = float(getattr(profile, "table_size_pt", self._caption_size_pt)) if is_formal else 9
        self._toc_size_pt = float(getattr(profile, "toc_size_pt", body_size)) if is_formal else body_size
        line_spacing = float(getattr(profile, "line_spacing", 1.5)) if is_formal else 1.45
        first_indent_cm = (
            float(getattr(profile, "first_line_indent_chars", 2.0)) * 0.37
            if is_formal
            else 0.74
        )

        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(margin_top)
        section.bottom_margin = Cm(margin_bottom)
        section.left_margin = Cm(margin_left)
        section.right_margin = Cm(margin_right)
        if is_formal:
            section.different_first_page_header_footer = True

        styles = doc.styles
        self._configure_paragraph_style(
            styles["Normal"], body_size, bold=False, color=self._paper_text_color,
            line_spacing=line_spacing, space_after=6,
        )
        self._configure_paragraph_style(
            styles["Body Text"], body_size, bold=False, color=self._paper_text_color,
            line_spacing=line_spacing, space_after=5,
        )
        if is_formal:
            styles["Body Text"].paragraph_format.first_line_indent = Cm(first_indent_cm)
            styles["Body Text"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        title_size = self._title_size_pt
        subtitle_size = self._subtitle_size_pt
        self._configure_paragraph_style(
            styles["Title"], title_size, bold=True, color=self._paper_text_color,
            line_spacing=1.1, space_after=12,
        )
        self._configure_paragraph_style(
            styles["Subtitle"], subtitle_size, bold=False,
            color=self._paper_accent_color,
            line_spacing=1.2, space_after=8,
        )
        heading_sizes = (
            (
                self._heading1_size_pt,
                self._heading2_size_pt,
                self._heading3_size_pt,
            )
            if is_formal
            else (18, 14.5, 12)
        )
        for (style_name, size, before, after) in zip(
            ("Heading 1", "Heading 2", "Heading 3"),
            heading_sizes,
            (20, 13, 8),
            (9, 6, 4),
        ):
            self._configure_paragraph_style(
                styles[style_name], size, bold=True, color=self._paper_accent_color,
                line_spacing=1.2, space_before=before, space_after=after,
            )
            styles[style_name].paragraph_format.keep_with_next = True

        caption_size = self._caption_size_pt
        self._configure_paragraph_style(
            styles["Caption"], caption_size, bold=False, color=self._paper_text_color,
            line_spacing=1.1, space_before=3, space_after=3,
        )
        if is_formal:
            styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        artifact_source = styles.add_style("Artifact Source", 1)
        self._configure_paragraph_style(
            artifact_source, 8, bold=False, color=self._paper_muted_color,
            line_spacing=1.0, space_before=0, space_after=8,
        )
        figure_lead = styles.add_style("Figure Lead", 1)
        self._configure_paragraph_style(
            figure_lead, body_size, bold=False, color=self._paper_muted_color,
            line_spacing=1.15, space_before=8, space_after=4,
        )
        figure_lead.paragraph_format.keep_with_next = True
        figure_note = styles.add_style("Figure Note", 1)
        self._configure_paragraph_style(
            figure_note, caption_size, bold=False, color=self._paper_muted_color,
            line_spacing=1.1, space_before=0, space_after=4,
        )
        figure_note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_note = styles.add_style("Table Note", 1)
        self._configure_paragraph_style(
            table_note, caption_size, bold=False, color=self._paper_muted_color,
            line_spacing=1.1, space_before=2, space_after=5,
        )
        self._configure_paragraph_style(
            styles["List Number"], body_size, bold=False, color=self._paper_text_color,
            line_spacing=line_spacing, space_after=4,
        )
        self._configure_paragraph_style(
            styles["List Bullet"], body_size, bold=False, color=self._paper_text_color,
            line_spacing=line_spacing, space_after=4,
        )
        reference_style = styles.add_style("Academic Reference", 1)
        self._configure_paragraph_style(
            reference_style, body_size, bold=False, color=self._paper_text_color,
            line_spacing=1.25, space_after=5,
        )
        reference_style.paragraph_format.left_indent = Cm(0.74)
        reference_style.paragraph_format.first_line_indent = Cm(-0.74)

        header = section.header.paragraphs[0]
        header.style = "Header"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.text = "学术论文" if is_formal else (project_name or "实验报告")
        for run in header.runs:
            self._set_run_font(run, 8, self._paper_muted_color)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.style = "Footer"
        footer.text = ""
        if not is_formal:
            footer.add_run("实验报告助手  ·  ")
        self._add_page_number_field(footer)
        for run in footer.runs:
            self._set_run_font(run, 8, self._paper_muted_color)

    def _configure_paragraph_style(
        self, style, size: float, *, bold: bool, color: str,
        line_spacing: float, space_before: float = 0, space_after: float = 0,
    ) -> None:
        style.font.name = getattr(self, "_body_font_en", BODY_FONT_EN)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = self._rgb(color)
        style._element.get_or_add_rPr().rFonts.set(
            qn("w:eastAsia"), getattr(self, "_body_font_cn", BODY_FONT_CN)
        )
        fmt = style.paragraph_format
        fmt.line_spacing = line_spacing
        fmt.space_before = Pt(space_before)
        fmt.space_after = Pt(space_after)

    def _set_run_font(self, run, size: float, color: str,
                      *, bold: bool = False) -> None:
        run.font.name = getattr(self, "_body_font_en", BODY_FONT_EN)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = self._rgb(color)
        run._element.get_or_add_rPr().rFonts.set(
            qn("w:eastAsia"), getattr(self, "_body_font_cn", BODY_FONT_CN)
        )

    @staticmethod
    def _rgb(hex_value: str):
        from docx.shared import RGBColor
        return RGBColor.from_string(hex_value)

    @staticmethod
    def _add_page_number_field(paragraph) -> None:
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    @staticmethod
    def _set_paragraph_bottom_border(paragraph, color: str = ACCENT_HEX) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "18")
        bottom.set(qn("w:space"), "6")
        bottom.set(qn("w:color"), color)
        borders.append(bottom)

    @staticmethod
    def _display_artifact_name(name: str) -> str:
        stem = Path(name or "未命名产物").stem.replace("_", " ")
        return stem or "未命名图表"

    def _image_width_inches(
        self, path: Path, doc: Document, layout_profile: str | None = None,
    ) -> float:
        """按正文宽度、图形 profile 和原始宽高比计算图片宽度。"""
        section = doc.sections[0]
        available = (
            section.page_width - section.left_margin - section.right_margin
        ) / 914400
        profile = str(layout_profile or "").strip().lower()
        if profile in {"half", "half_width", "side_by_side", "paired", "two_column"}:
            return round(min(float(available) / 2 - 0.12, 3.05), 2)
        try:
            from PIL import Image
            with Image.open(path) as image:
                width, height = image.size
                ratio = width / max(height, 1)
                # 过窄的图保持适度可读宽度；宽图最多占正文宽度。
                if ratio < 0.9:
                    return round(min(float(available), 5.6), 2)
        except Exception:
            pass
        return round(min(float(available), 6.6), 2)

    def _add_artifact_error(self, doc: Document, message: str) -> None:
        style = "Body Text" if getattr(self, "_reader_first_paper", False) else "Artifact Source"
        paragraph = doc.add_paragraph(style=style)
        if getattr(self, "_reader_first_paper", False):
            paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.add_run(f"[{message}]")

    def _render_cover(
        self, doc: Document, project_name: str,
        project_topic: str, plan=None,
    ) -> None:
        """渲染封面。"""
        if getattr(self, "_formal_paper", False):
            self._render_formal_cover(doc, project_name, project_topic, plan)
            return
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(78)

        eyebrow = doc.add_paragraph()
        eyebrow.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = eyebrow.add_run("实验报告助手  /  数据分析实验")
        self._set_run_font(run, 10, ACCENT_HEX, bold=True)

        heading = doc.add_paragraph(style="Title")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(project_topic or "实验报告")
        self._set_run_font(run, 26, TEXT_HEX, bold=True)
        self._set_paragraph_bottom_border(heading)

        p = doc.add_paragraph(style="Subtitle")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"项目：{project_name}")
        self._set_run_font(run, 14, ACCENT_HEX)

        meta = doc.add_paragraph(style="Artifact Source")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.paragraph_format.space_before = Pt(18)
        meta.paragraph_format.space_after = Pt(4)
        meta.add_run("实验报告  ·  ").bold = True
        meta.add_run(
            f"生成日期：{datetime.now(timezone.utc).strftime('%Y-%m-%d')}"
        )
        for meta_run in meta.runs:
            self._set_run_font(meta_run, 9.5, MUTED_HEX, bold=meta_run.bold)

        doc.add_page_break()

    def _render_formal_cover(
        self, doc: Document, project_name: str, project_topic: str, plan,
    ) -> None:
        """渲染正式论文封面，突出题名、研究类型和数据边界。"""
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(72)

        label = doc.add_paragraph()
        label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = label.add_run("数据分析论文解读与教学性复核")
        self._set_run_font(run, 11, MUTED_HEX, bold=True)

        title = doc.add_paragraph(style="Title")
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(30)
        title.paragraph_format.space_after = Pt(16)
        title_run = title.add_run(plan.formal_title or project_topic or "学术论文解读")
        self._set_run_font(title_run, 22, TEXT_HEX, bold=True)

        subtitle = doc.add_paragraph(style="Subtitle")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(plan.formal_subtitle or project_name)
        self._set_run_font(subtitle_run, 13, self._paper_accent_color)

        blocked_metadata_keys = {
            "执行批次", "execution_run_id", "run_id", "sha256",
            "artifact_id", "file_path", "json path", "json 路径",
            "作者", "author", "单位", "affiliation", "机构",
            "提交日期", "date",
        }
        metadata = {
            str(key): str(value)
            for key, value in plan.formal_metadata
            if str(key).strip().lower() not in blocked_metadata_keys
        }
        metadata_lookup = {
            str(key).strip().lower(): str(value).strip()
            for key, value in plan.formal_metadata
        }
        author = (
            metadata_lookup.get("作者")
            or metadata_lookup.get("author")
            or metadata_lookup.get("作者姓名")
        )
        affiliation = (
            metadata_lookup.get("单位")
            or metadata_lookup.get("affiliation")
            or metadata_lookup.get("机构")
        )
        submission_date = (
            metadata_lookup.get("提交日期")
            or metadata_lookup.get("date")
            or datetime.now(timezone.utc).strftime("%Y-%m-%d")
        )
        metadata.setdefault("研究类型", "公开论文解读与本地描述性复核")
        metadata.setdefault("数据边界", "原始公开 CSV 与论文报告样本分开陈述")
        metadata.setdefault("生成日期", submission_date)

        if author or affiliation:
            identity = doc.add_paragraph()
            identity.alignment = WD_ALIGN_PARAGRAPH.CENTER
            identity.paragraph_format.space_before = Pt(18)
            identity.paragraph_format.space_after = Pt(2)
            if author:
                author_run = identity.add_run(f"作者：{author}")
                self._set_run_font(author_run, 11.5, TEXT_HEX, bold=True)
            if affiliation:
                affiliation_paragraph = doc.add_paragraph()
                affiliation_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                affiliation_paragraph.paragraph_format.space_after = Pt(2)
                affiliation_run = affiliation_paragraph.add_run(f"单位：{affiliation}")
                self._set_run_font(affiliation_run, 10.5, MUTED_HEX)
            date_paragraph = doc.add_paragraph()
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_paragraph.paragraph_format.space_after = Pt(16)
            date_run = date_paragraph.add_run(f"提交日期：{submission_date}")
            self._set_run_font(date_run, 9.5, MUTED_HEX)

        rule = doc.add_paragraph()
        rule.paragraph_format.space_before = Pt(8)
        rule.paragraph_format.space_after = Pt(12)
        self._set_paragraph_bottom_border(rule, self._paper_accent_color)

        info_heading = doc.add_paragraph()
        info_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_heading.paragraph_format.space_after = Pt(5)
        info_run = info_heading.add_run("研究信息")
        self._set_run_font(info_run, 9.5, self._paper_accent_color, bold=True)
        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_widths(table, (1.6, 4.7))
        for key, value in metadata.items():
            cells = table.add_row().cells
            cells[0].text = str(key)
            cells[1].text = str(value)
            for index, cell in enumerate(cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(cell, top=60, start=120, bottom=60, end=120)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        self._set_run_font(
                            run, self._table_size_pt, self._paper_accent_color if index == 0 else self._paper_text_color,
                            bold=index == 0,
                        )
        self._set_formal_table_borders(table)
        doc.add_page_break()

    def _enforce_formal_sufficiency(self, document_plan) -> None:
        """formal 交付必须通过内容充分性门禁，避免生成空壳论文。"""
        if not getattr(self, "_formal_paper", False):
            return
        report = getattr(document_plan, "sufficiency", None)
        if report is None or report.publishable:
            return
        codes = tuple(
            issue.code
            for issue in getattr(report, "issues", ())
            if getattr(issue, "code", "")
        )
        error = AppError(
            code="WORD_CONTENT_NOT_PUBLISHABLE",
            message="formal Word 交付被内容充分性门禁阻断：" + ", ".join(codes),
            field="sufficiency",
        )
        # AppError 的公共合同保持兼容；codes 作为结构化门禁元数据暴露。
        error.codes = codes
        raise error

    def _add_formal_section(self, doc: Document, *, page_format: str) -> None:
        """创建 formal 前置/正文节，并显式设置页码格式。"""
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        profile = getattr(self, "_publication_profile", None)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(float(getattr(profile, "page_margin_top_cm", PAGE_MARGIN_TOP_CM)))
        section.bottom_margin = Cm(float(getattr(profile, "page_margin_bottom_cm", PAGE_MARGIN_BOTTOM_CM)))
        section.left_margin = Cm(float(getattr(profile, "page_margin_left_cm", PAGE_MARGIN_LEFT_CM)))
        section.right_margin = Cm(float(getattr(profile, "page_margin_right_cm", PAGE_MARGIN_RIGHT_CM)))
        section.different_first_page_header_footer = False
        self._set_section_page_number_format(section, page_format)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.text = "学术论文"
        for run in header.runs:
            self._set_run_font(run, 8, MUTED_HEX)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = ""
        self._add_page_number_field(footer)
        for run in footer.runs:
            self._set_run_font(run, 8, MUTED_HEX)

    @staticmethod
    def _set_section_page_number_format(section, page_format: str) -> None:
        sect_pr = section._sectPr
        pg_num_type = sect_pr.find(qn("w:pgNumType"))
        if pg_num_type is None:
            pg_num_type = OxmlElement("w:pgNumType")
            sect_pr.append(pg_num_type)
        pg_num_type.set(qn("w:fmt"), page_format)
        if page_format in {"decimal", "lowerRoman"}:
            pg_num_type.set(qn("w:start"), "1")

    @staticmethod
    def _add_field(paragraph, instruction: str, result: str = ""):
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        begin.set(qn("w:dirty"), "true")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" {instruction} "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        result_node = OxmlElement("w:t")
        result_node.text = result
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, instr, separate, result_node, end))
        return run

    def _add_paragraph_bookmark(self, paragraph, bookmark_name: str) -> None:
        """给目录可回指的章节标题添加书签。"""
        self._bookmark_id += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(self._bookmark_id))
        start.set(qn("w:name"), bookmark_name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(self._bookmark_id))
        paragraph._p.insert(0, start)
        paragraph._p.append(end)

    def _add_bookmarked_sequence(
        self, paragraph, bookmark_name: str, sequence_name: str, visible_label: str,
    ) -> None:
        visible_label = str(visible_label).strip()
        match = re.match(r"^(图|表)\s+(.+)$", visible_label)
        prefix = f"{match.group(1)} " if match else ""
        number = match.group(2) if match else visible_label
        self._bookmark_id += 1
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(self._bookmark_id))
        start.set(qn("w:name"), bookmark_name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(self._bookmark_id))
        paragraph._p.append(start)
        paragraph.add_run(prefix)
        self._add_field(paragraph, f"SEQ {sequence_name} \\* ARABIC", number)
        paragraph._p.append(end)

    @classmethod
    def _add_cross_reference(cls, paragraph, bookmark_name: str) -> None:
        cls._add_field(paragraph, f"REF {bookmark_name}", "1")
        paragraph.add_run("（第")
        cls._add_field(paragraph, f"PAGEREF {bookmark_name}", "1")
        paragraph.add_run("页）")
    def _render_front_matter(self, doc: Document, plan) -> None:
        """渲染摘要、关键词和静态目录。"""

        if getattr(self, "_formal_paper", False):
            self._render_formal_front_matter(doc, plan)
            return

        doc.add_heading("摘要", level=1)
        abstract = doc.add_paragraph(plan.abstract, style="Body Text")
        abstract.paragraph_format.first_line_indent = Cm(0.74)

        keyword = doc.add_paragraph(style="Body Text")
        keyword.paragraph_format.first_line_indent = Cm(0.74)
        keyword.add_run("关键词：").bold = True
        keyword.add_run("；".join(plan.keywords))

        doc.add_page_break()
        doc.add_heading("目录", level=1)
        toc_note = doc.add_paragraph(style="Artifact Source")
        toc_note.add_run("目录依据当前已确认大纲生成；文档结构变更后需重新生成。").italic = True
        for chapter in plan.chapters:
            entry = doc.add_paragraph(style="Body Text")
            entry.paragraph_format.left_indent = Cm(0.4)
            entry.paragraph_format.space_after = Pt(3)
            entry.add_run(f"第 {chapter.number} 章  {chapter.title}")
            for section in chapter.sections:
                sub = doc.add_paragraph(style="Body Text")
                sub.paragraph_format.left_indent = Cm(1.1)
                sub.paragraph_format.space_after = Pt(2)
                sub.add_run(str(section.get("title", "内容")))
        doc.add_page_break()


    def _render_formal_front_heading(self, doc: Document, title: str, bookmark: str):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(12)
        run = paragraph.add_run(title)
        self._set_run_font(run, 16, self._paper_text_color, bold=True)
        self._set_paragraph_bottom_border(paragraph, self._paper_accent_color)
        self._add_paragraph_bookmark(paragraph, bookmark)
        return paragraph

    def _render_abstract_sections(
        self, doc: Document, sections: tuple[tuple[str, str], ...], *, english: bool = False,
    ) -> None:
        for label, content in sections:
            paragraph = doc.add_paragraph(style="Body Text")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.35
            paragraph.paragraph_format.space_after = Pt(8)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            separator = ": " if english else "："
            label_run = paragraph.add_run(f"{label}{separator}")
            self._set_run_font(label_run, self._body_size_pt, self._paper_text_color, bold=True)
            content_run = paragraph.add_run(content)
            self._set_run_font(content_run, self._body_size_pt, self._paper_text_color)

    def _render_formal_front_matter(self, doc: Document, plan) -> None:
        self._render_formal_front_heading(doc, "摘要", "front_abstract_cn")
        if plan.abstract_sections:
            self._render_abstract_sections(doc, plan.abstract_sections)
        else:
            abstract = doc.add_paragraph(style="Body Text")
            abstract.paragraph_format.first_line_indent = Cm(0.74)
            abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abstract.add_run(plan.abstract)

        keyword = doc.add_paragraph(style="Body Text")
        keyword.paragraph_format.first_line_indent = Cm(0)
        keyword.paragraph_format.space_before = Pt(4)
        keyword.paragraph_format.keep_together = True
        keyword_run = keyword.add_run("关键词：")
        self._set_run_font(keyword_run, self._body_size_pt, self._paper_text_color, bold=True)
        values_run = keyword.add_run("；".join(plan.keywords))
        self._set_run_font(values_run, self._body_size_pt, self._paper_text_color)

        doc.add_page_break()
        self._render_formal_front_heading(doc, "ABSTRACT", "front_abstract_en")
        if plan.abstract_sections_en:
            self._render_abstract_sections(doc, plan.abstract_sections_en, english=True)
        elif plan.abstract_en:
            abstract_en = doc.add_paragraph(style="Body Text")
            abstract_en.paragraph_format.first_line_indent = Cm(0.74)
            abstract_en.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            abstract_en.add_run(plan.abstract_en)
        keywords_en = doc.add_paragraph(style="Body Text")
        keywords_en.paragraph_format.first_line_indent = Cm(0)
        keywords_en.paragraph_format.space_before = Pt(4)
        keywords_en.paragraph_format.keep_together = True
        keyword_en_run = keywords_en.add_run("KEYWORDS: ")
        self._set_run_font(keyword_en_run, self._body_size_pt, self._paper_text_color, bold=True)
        values_en_run = keywords_en.add_run("; ".join(plan.keywords))
        self._set_run_font(values_en_run, self._body_size_pt, self._paper_text_color)

        doc.add_page_break()
        self._render_formal_front_heading(doc, "目录", "front_toc")

        def add_entry(label: str, bookmark: str, level: int) -> None:
            paragraph = doc.add_paragraph(style="Body Text")
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.left_indent = Cm(
                {1: 0.15, 2: 0.75, 3: 1.35}.get(level, 0.75)
            )
            paragraph.paragraph_format.space_before = Pt(4 if level == 1 else 0)
            paragraph.paragraph_format.space_after = Pt(1 if level == 1 else 0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS,
            )
            run = paragraph.add_run(f"{label}\t")
            self._set_run_font(
                run, self._toc_size_pt,
                self._paper_text_color,
                bold=level == 1,
            )
            page_run = self._add_field(paragraph, f"PAGEREF {bookmark}", "1")
            self._set_run_font(page_run, self._toc_size_pt, self._paper_text_color)

        for label, bookmark in (
            ("摘要", "front_abstract_cn"),
            ("ABSTRACT", "front_abstract_en"),
            ("目录", "front_toc"),
            ("图目录", "front_figures"),
            ("表目录", "front_tables"),
        ):
            add_entry(label, bookmark, 1)

        for chapter in plan.chapters:
            add_entry(
                f"第 {chapter.number} 章  {chapter.title}",
                f"chapter_{chapter.number}",
                1,
            )
            for section_index, section in enumerate(chapter.sections, start=1):
                add_entry(
                    f"{chapter.number}.{section_index}  {section.get('title', '内容')}",
                    f"section_{chapter.number}_{section_index}",
                    2,
                )
                for subsection_index, subsection in enumerate(
                    section.get("subsections", []) or [], start=1
                ):
                    subsection_title = (
                        subsection.get("title", "内容")
                        if isinstance(subsection, dict)
                        else str(subsection)
                    )
                    add_entry(
                        f"{chapter.number}.{section_index}.{subsection_index}  {subsection_title}",
                        f"section_{chapter.number}_{section_index}_{subsection_index}",
                        3,
                    )

    def _render_formal_figure_table_catalog(
        self, doc: Document, artifacts: list[dict], plan=None,
    ) -> None:
        figure_entries: list[tuple[str, int, str, str, str]] = []
        table_entries: list[tuple[str, int, str, str, str]] = []
        figure_index = 0
        table_index = 0

        def add_artifact_entry(
            artifact_type: str, caption: str, chapter_label: str,
        ) -> None:
            nonlocal figure_index, table_index
            if artifact_type == "CHART_PNG":
                figure_index += 1
                figure_entries.append(
                    ("图", figure_index, caption, f"fig_{figure_index}", chapter_label)
                )
            elif artifact_type == "TABLE_CSV":
                table_index += 1
                table_entries.append(
                    ("表", table_index, caption, f"tbl_{table_index}", chapter_label)
                )

        if plan is not None:
            for chapter in plan.chapters:
                chapter_label = f"第 {chapter.number} 章  {chapter.title}"
                for section in chapter.sections:
                    layout = plan_section_layouts([section], artifacts)[0]
                    if layout.layout_kind == LayoutKind.DATA_OVERVIEW and layout.metrics:
                        add_artifact_entry("TABLE_CSV", "数据概览", chapter_label)

                    source_type = str(section.get("source_type", ""))
                    if source_type != "EXECUTION" and not section.get("artifact_group"):
                        continue
                    source_ids = section.get("source_ids", []) or []
                    relevant = [
                        art for art in artifacts
                        if (
                            section.get("artifact_group")
                            and art.get("artifact_group") == section.get("artifact_group")
                        )
                        or (
                            not section.get("artifact_group")
                            and (not source_ids or art.get("execution_run_id") in source_ids)
                        )
                    ]
                    for artifact in relevant:
                        add_artifact_entry(
                            str(artifact.get("artifact_type", "")),
                            str(
                                artifact.get("figure_caption")
                                or artifact.get("table_caption")
                                or "未命名图表"
                            ).strip(),
                            chapter_label,
                        )
        else:
            for artifact in artifacts:
                add_artifact_entry(
                    str(artifact.get("artifact_type", "")),
                    str(
                        artifact.get("figure_caption")
                        or artifact.get("table_caption")
                        or "未命名图表"
                    ).strip(),
                    "图表",
                )

        if not figure_entries and not table_entries:
            return

        doc.add_page_break()

        def render_catalog(
            title: str, bookmark: str, entries: list[tuple[str, int, str, str, str]],
        ) -> None:
            self._render_formal_front_heading(doc, title, bookmark)
            last_chapter = None
            for prefix, number, caption, entry_bookmark, chapter_label in entries:
                if chapter_label != last_chapter:
                    chapter_paragraph = doc.add_paragraph(style="Body Text")
                    chapter_paragraph.paragraph_format.first_line_indent = Cm(0)
                    chapter_paragraph.paragraph_format.left_indent = Cm(0.15)
                    chapter_paragraph.paragraph_format.space_before = Pt(4)
                    chapter_paragraph.paragraph_format.space_after = Pt(1)
                    chapter_run = chapter_paragraph.add_run(chapter_label)
                    self._set_run_font(
                        chapter_run, self._toc_size_pt, self._paper_muted_color, bold=True,
                    )
                    last_chapter = chapter_label
                paragraph = doc.add_paragraph(style="Body Text")
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.left_indent = Cm(0.65)
                paragraph.paragraph_format.space_after = Pt(1)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.tab_stops.add_tab_stop(
                    Inches(6.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS,
                )
                ref_run = self._add_field(paragraph, f"REF {entry_bookmark}", f"{prefix} {number}")
                self._set_run_font(ref_run, self._toc_size_pt, self._paper_text_color)
                caption_run = paragraph.add_run(f"  {caption}\t")
                self._set_run_font(caption_run, self._toc_size_pt, self._paper_text_color)
                page_run = self._add_field(paragraph, f"PAGEREF {entry_bookmark}", "1")
                self._set_run_font(page_run, self._toc_size_pt, self._paper_text_color)

        if figure_entries:
            render_catalog("图目录", "front_figures", figure_entries)
        if table_entries:
            render_catalog("表目录", "front_tables", table_entries)

    def _render_chapter(
        self,
        doc: Document,
        chapter: ThesisChapterPlan,
        artifacts: list[dict],
    ) -> None:
        """渲染一个论文一级章节。"""

        heading = doc.add_heading(
            f"第 {chapter.number} 章  {chapter.title}", level=1,
        )
        if getattr(self, "_formal_paper", False):
            self._add_paragraph_bookmark(heading, f"chapter_{chapter.number}")
        # 论文正文采用连续版式；仅由分页引擎在确有空间不足时换页，
        # 避免每个一级章节强制制造半页空白。
        heading.paragraph_format.page_break_before = False
        self._chapter_figure_counts[chapter.number] = 0
        self._chapter_table_counts[chapter.number] = 0
        for section_index, section in enumerate(chapter.sections, start=1):
            plan = plan_section_layouts([section], artifacts)[0]
            self._render_section(
                doc, section, artifacts, layout_plan=plan,
                chapter_number=chapter.number,
                section_number=section_index,
            )

    def _render_references(self, doc: Document, references: tuple[str, ...]) -> None:
        """渲染参考资料；没有外部资料时显式说明。"""

        doc.add_page_break()
        doc.add_heading(
            "参考文献" if getattr(self, "_formal_paper", False) else "参考资料",
            level=1,
        )
        for reference in references:
            paragraph = doc.add_paragraph(
                style=("Academic Reference" if getattr(self, "_formal_paper", False)
                       else "List Number")
            )
            paragraph.add_run(reference)

    def _render_section(
        self,
        doc: Document,
        section: dict,
        artifacts: list[dict],
        *,
        layout_plan: SectionLayoutPlan | None = None,
        chapter_number: int | None = None,
        section_number: int | None = None,
    ) -> None:
        """渲染单个章节；默认渲染使用 SPEC 0033 语义版式计划。"""
        title = section.get("title", "")
        content = section.get("content", "")
        source_type = section.get("source_type", "")
        source_ids = section.get("source_ids", []) or []
        kind = layout_plan.layout_kind if layout_plan else LayoutKind.NARRATIVE

        # 章节标题
        heading_title = title
        if getattr(self, "_formal_paper", False) and chapter_number and section_number:
            heading_title = f"{chapter_number}.{section_number}  {title}"
        heading = doc.add_heading(heading_title, level=2 if chapter_number else 1)
        if getattr(self, "_formal_paper", False) and chapter_number and section_number:
            self._add_paragraph_bookmark(
                heading, f"section_{chapter_number}_{section_number}",
            )
        heading.paragraph_format.page_break_before = False

        # 章节内容：方法章节按步骤分段，数据概览用轻量指标表，其余按论文
        # 段落计划渲染，避免把“说明、解释和边界”压成一个长段落。
        if kind == LayoutKind.METHOD_FLOW and layout_plan and layout_plan.steps:
            self._render_method_steps(doc, layout_plan.steps, source_ids=source_ids)
        elif kind == LayoutKind.DATA_OVERVIEW and layout_plan and layout_plan.metrics:
            self._render_data_overview(
                doc,
                layout_plan.metrics,
                content,
                source_ids=source_ids,
                chapter_number=chapter_number,
                caption_text=str(section.get("table_caption") or "数据概览"),
            )
        elif content:
            paragraphs = section.get("paragraphs") or [content]
            for paragraph_text in paragraphs:
                if str(paragraph_text).strip():
                    paragraph = doc.add_paragraph(style="Body Text")
                    paragraph.add_run(str(paragraph_text).strip())
                    self._append_citations(paragraph, source_ids)


        for subsection_index, subsection in enumerate(
            section.get("subsections", []) or [], start=1
        ):
            if not isinstance(subsection, dict):
                continue
            subsection_title = str(subsection.get("title", "内容")).strip()
            subsection_heading = doc.add_heading(subsection_title, level=3)
            if (
                getattr(self, "_formal_paper", False)
                and chapter_number
                and section_number
            ):
                self._add_paragraph_bookmark(
                    subsection_heading,
                    f"section_{chapter_number}_{section_number}_{subsection_index}",
                )
            subsection_content = str(subsection.get("content", "")).strip()
            subsection_paragraphs = subsection.get("paragraphs") or (
                [subsection_content] if subsection_content else []
            )
            for subsection_text in subsection_paragraphs:
                if str(subsection_text).strip():
                    subsection_paragraph = doc.add_paragraph(style="Body Text")
                    subsection_paragraph.add_run(str(subsection_text).strip())
                    self._append_citations(subsection_paragraph, source_ids)

        figure_lead = str(section.get("figure_lead", "")).strip()
        if (
            figure_lead
            and (source_type == "EXECUTION" or section.get("artifact_group"))
            and (
                not getattr(self, "_reader_first_paper", False)
                or section.get("reader_figure_lead")
            )
        ):
            lead_style = "Body Text" if getattr(self, "_formal_paper", False) else "Figure Lead"
            lead = doc.add_paragraph(style=lead_style)
            if getattr(self, "_formal_paper", False):
                lead.paragraph_format.first_line_indent = Cm(0)
                lead.paragraph_format.keep_with_next = True
            lead.add_run(figure_lead)

        figure_takeaway = str(section.get("figure_takeaway", "")).strip()

        # 执行产物或显式逻辑图引用。逻辑图可以挂在 EVIDENCE/ANALYSIS
        # 章节，但仍必须通过 artifact_group 与已确认大纲关联。
        if source_type == "EXECUTION" or section.get("artifact_group"):
            self._render_artifacts(
                doc,
                artifacts,
                source_ids,
                artifact_group=section.get("artifact_group"),
                chapter_number=chapter_number,
                figure_takeaway=figure_takeaway,
            )
        if figure_takeaway and (source_type == "EXECUTION" or section.get("artifact_group")):
            takeaway = doc.add_paragraph(style="Body Text")
            takeaway.paragraph_format.keep_together = True
            takeaway.paragraph_format.space_before = Pt(4)
            takeaway.add_run(figure_takeaway)
            self._append_citations(takeaway, source_ids)

    def _render_method_steps(
        self, doc: Document, steps: tuple[str, ...], *, source_ids: list[str] | None = None,
    ) -> None:
        """以论文正文节奏呈现分析步骤，避免把方法压成一堵文字。"""

        for step in steps:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Cm(0.74)
            paragraph.paragraph_format.first_line_indent = Cm(-0.37)
            run = paragraph.add_run(step)
            self._set_run_font(run, getattr(self, "_body_size_pt", 12), TEXT_HEX)
        if source_ids:
            if getattr(self, "_reader_first_paper", False):
                note = doc.add_paragraph(style="Body Text")
                note.paragraph_format.first_line_indent = Cm(0)
                note.add_run("方法依据见文献")
            else:
                note = doc.add_paragraph(style="Artifact Source")
                note.add_run("方法依据：")
            self._append_citations(note, source_ids)

    def _render_data_overview(
        self,
        doc: Document,
        metrics: tuple[tuple[str, str], ...],
        content: str,
        *,
        source_ids: list[str] | None = None,
        chapter_number: int | None = None,
        caption_text: str | None = None,
    ) -> None:
        """以紧凑二列表呈现真实数据概览，再保留原始说明。"""

        self._table_index += 1
        table_label = f"表 {self._table_index}"
        caption = doc.add_paragraph(style="Caption")
        caption.paragraph_format.keep_with_next = True
        if getattr(self, "_formal_paper", False):
            self._add_bookmarked_sequence(
                caption, f"tbl_{self._table_index}", "Table", table_label,
            )
            caption.add_run(f"  {caption_text or '数据概览'}").bold = True
        else:
            caption.add_run(f"{table_label}  {caption_text or '数据概览'}").bold = True

        table = doc.add_table(rows=1, cols=2)
        self._mark_table_header(table.rows[0])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        self._set_table_widths(table, (2.1, 4.2))
        header = table.rows[0].cells
        header[0].text = "指标"
        header[1].text = "值"
        for cell in header:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if not getattr(self, "_formal_paper", False):
                self._set_cell_shading(cell, TABLE_HEADER_HEX)
            self._set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    self._set_run_font(
                        run, 9, TEXT_HEX if getattr(self, "_formal_paper", False) else "FFFFFF",
                        bold=True,
                    )

        for label, value in metrics:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
            for index, cell in enumerate(cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
                if not getattr(self, "_formal_paper", False):
                    self._set_cell_shading(cell, TABLE_ALT_HEX if len(table.rows) % 2 == 0 else "FFFFFF")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index else WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        self._set_run_font(run, getattr(self, "_table_size_pt", 9), TEXT_HEX, bold=index == 0)
        for row in table.rows:
            self._mark_table_row_no_split(row)
        if getattr(self, "_formal_paper", False):
            self._set_formal_table_borders(table)
        else:
            self._set_table_borders(table)

        if content:
            paragraph = doc.add_paragraph(style="Body Text")
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.add_run(content)
            self._append_citations(paragraph, source_ids or [])

    def _render_artifacts(
        self, doc: Document, artifacts: list[dict],
        source_ids: list[str],
        artifact_group: str | None = None,
        *, chapter_number: int | None = None,
        figure_takeaway: str | None = None,
    ) -> None:
        """渲染执行产物；显式 side_by_side profile 才启用双图并排。"""
        relevant = []
        for art in artifacts:
            if artifact_group and art.get("artifact_group") != artifact_group:
                continue
            if artifact_group:
                relevant.append(art)
                continue
            if not source_ids or art.get("execution_run_id") in source_ids:
                relevant.append(art)

        if not relevant:
            return

        if not getattr(self, "_formal_paper", False):
            doc.add_heading("执行产物", level=2)

        index = 0
        while index < len(relevant):
            if (
                index + 1 < len(relevant)
                and self._should_render_side_by_side(
                    relevant[index], relevant[index + 1],
                )
            ):
                self._render_side_by_side_charts(
                    doc,
                    relevant[index:index + 2],
                    keep_last_note_with_next=bool(
                        figure_takeaway and index + 2 >= len(relevant)
                    ),
                )
                index += 2
                continue
            self._render_single_artifact(
                doc,
                relevant[index],
                chapter_number=chapter_number,
                keep_note_with_next=bool(
                    figure_takeaway and index == len(relevant) - 1
                ),
            )
            index += 1

    @staticmethod
    def _should_render_side_by_side(first: dict, second: dict) -> bool:
        allowed = {"side_by_side", "paired", "two_column"}
        return (
            first.get("artifact_type") == "CHART_PNG"
            and second.get("artifact_type") == "CHART_PNG"
            and str(first.get("file_path", "")).strip()
            and str(second.get("file_path", "")).strip()
            and str(first.get("figure_layout_profile", "")).strip().lower() in allowed
            and str(second.get("figure_layout_profile", "")).strip().lower() in allowed
            and Path(first["file_path"]).exists()
            and Path(second["file_path"]).exists()
        )

    def _render_side_by_side_charts(
        self,
        doc: Document,
        charts: list[dict],
        *,
        keep_last_note_with_next: bool = False,
    ) -> None:
        """把两个已声明成 panel pair 的真实图表排成双栏。"""
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_before = Pt(8)
        image_paragraph.paragraph_format.space_after = Pt(2)
        image_paragraph.paragraph_format.keep_with_next = True

        figure_numbers: list[int] = []
        for offset, art in enumerate(charts):
            path = Path(str(art["file_path"]))
            self._figure_index += 1
            figure_numbers.append(self._figure_index)
            run = image_paragraph.add_run()
            run.add_picture(
                str(path),
                width=Inches(
                    self._image_width_inches(
                        path, doc, "side_by_side",
                    )
                ),
            )
            if offset < len(charts) - 1:
                image_paragraph.add_run("    ")

        for art, figure_number in zip(charts, figure_numbers):
            caption = doc.add_paragraph(style="Caption")
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.keep_with_next = True
            label = f"图 {figure_number}"
            if getattr(self, "_formal_paper", False):
                self._add_bookmarked_sequence(
                    caption, f"fig_{figure_number}", "Figure", label,
                )
                caption.add_run(
                    f"  {art.get('figure_caption') or self._display_artifact_name(art.get('name', ''))}"
                ).bold = True
            else:
                caption.add_run(
                    f"{label}  {art.get('figure_caption') or self._display_artifact_name(art.get('name', ''))}"
                ).bold = True

            note_text = str(art.get("figure_note") or "").strip()
            if not note_text:
                continue
            if getattr(self, "_reader_first_paper", False):
                note = doc.add_paragraph(style="Figure Note")
            else:
                note = doc.add_paragraph(style="Artifact Source")
            note.paragraph_format.keep_with_next = (
                keep_last_note_with_next and offset == len(charts) - 1
            )
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note.add_run(note_text)

    def _render_csv_table(
        self,
        doc: Document,
        csv_path: Path,
        *,
        chapter_number: int | None = None,
        caption_text: str | None = None,
        keep_note_with_next: bool = False,
    ) -> None:
        """将 CSV 文件渲染为 Word 表格（前 10 行，前 6 列）。"""
        import csv

        with open(csv_path, "r", encoding="utf-8", newline="") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return

        # 限制行数和列数，避免原始数据把正文撑坏。
        rows = rows[:MAX_TABLE_ROWS]
        rows = [row[:MAX_TABLE_COLS] for row in rows]

        self._table_index += 1
        table_label = f"表 {self._table_index}"
        caption = doc.add_paragraph(style="Caption")
        caption.paragraph_format.keep_with_next = True
        if getattr(self, "_formal_paper", False):
            self._add_bookmarked_sequence(
                caption, f"tbl_{self._table_index}", "Table", table_label,
            )
            caption.add_run(
                f"  {caption_text or ('数据表：' + self._display_artifact_name(csv_path.name))}"
            ).bold = True
        else:
            caption.add_run(
                f"{table_label}  {caption_text or ('数据表：' + self._display_artifact_name(csv_path.name))}"
            ).bold = True

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        self._mark_table_header(table.rows[0])
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        column_width = 6.3 / max(1, len(rows[0]))
        self._set_table_widths(table, tuple(column_width for _ in rows[0]))
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j >= len(table.rows[i].cells):
                    continue
                cell = table.rows[i].cells[j]
                cell.text = str(cell_text)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(cell, top=65, start=90, bottom=65, end=90)
                if not getattr(self, "_formal_paper", False):
                    self._set_cell_shading(
                        cell,
                        TABLE_HEADER_HEX if i == 0 else (TABLE_ALT_HEX if i % 2 == 0 else "FFFFFF"),
                    )
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        self._set_run_font(run, getattr(self, "_table_size_pt", 9), "FFFFFF" if i == 0 else TEXT_HEX,
                                           bold=i == 0)
        for row in table.rows:
            self._mark_table_row_no_split(row)
        if getattr(self, "_formal_paper", False):
            self._set_formal_table_borders(table)
        else:
            self._set_table_borders(table)
        if len(rows) >= MAX_TABLE_ROWS or any(len(row) >= MAX_TABLE_COLS for row in rows):
            note_style = "Table Note" if getattr(self, "_reader_first_paper", False) else "Artifact Source"
            note = doc.add_paragraph(style=note_style)
            note.paragraph_format.keep_with_next = keep_note_with_next
            note.add_run(f"注：表格展示前 {MAX_TABLE_ROWS - 1} 行数据和前 {MAX_TABLE_COLS} 列字段。")

    def _render_appendix(self, doc: Document,
                          artifacts: list[dict]) -> None:
        """渲染附录：执行产物索引。"""
        if not artifacts or (
            getattr(self, "_formal_paper", False)
            and not getattr(self, "_include_audit_appendix", False)
        ):
            return

        doc.add_page_break()
        doc.add_heading("附录：执行产物索引", level=1)
        for art in artifacts:
            name = art.get("name", "")
            art_type = art.get("artifact_type", "")
            run_id = art.get("execution_run_id", "")
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{name}（{art_type}）")
            if run_id:
                paragraph.add_run(f"  ·  执行批次：{run_id}")

    @staticmethod
    def _set_cell_shading(cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), fill)

    @staticmethod
    def _set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        margins = tc_pr.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            tc_pr.append(margins)
        for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
            node = margins.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    @staticmethod
    def _set_table_borders(table) -> None:
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge, color, size in (
            ("top", TABLE_HEADER_HEX, "12"),
            ("bottom", TABLE_HEADER_HEX, "12"),
            ("insideH", "C9D9DE", "4"),
            ("insideV", "FFFFFF", "0"),
            ("left", "FFFFFF", "0"),
            ("right", "FFFFFF", "0"),
        ):
            tag = qn(f"w:{edge}")
            border = borders.find(tag)
            if border is None:
                border = OxmlElement(f"w:{edge}")
                borders.append(border)
            border.set(qn("w:val"), "single" if size != "0" else "nil")
            border.set(qn("w:sz"), size)
            border.set(qn("w:color"), color)

    @staticmethod
    def _set_formal_table_borders(table) -> None:
        """学术论文三线表：仅保留顶线、表头线和底线。"""
        tbl_pr = table._tbl.tblPr
        borders = tbl_pr.first_child_found_in("w:tblBorders")
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tbl_pr.append(borders)
        for edge, color, size, value in (
            ("top", "222222", "12", "single"),
            ("bottom", "222222", "12", "single"),
            ("insideH", "9CA3AF", "4", "single"),
            ("insideV", "FFFFFF", "0", "nil"),
            ("left", "FFFFFF", "0", "nil"),
            ("right", "FFFFFF", "0", "nil"),
        ):
            tag = qn(f"w:{edge}")
            border = borders.find(tag)
            if border is None:
                border = OxmlElement(f"w:{edge}")
                borders.append(border)
            border.set(qn("w:val"), value)
            border.set(qn("w:sz"), size)
            border.set(qn("w:color"), color)

    def _append_citations(self, paragraph, source_ids: list[str]) -> None:
        """在正文段末追加已规划的顺序编码引文。"""
        if not getattr(self, "_formal_paper", False):
            return
        citations = [
            self._citation_map[source_id]
            for source_id in source_ids
            if source_id in self._citation_map
        ]
        if not citations:
            return
        citations = sorted(set(citations))
        run = paragraph.add_run(" " + "[" + ", ".join(str(number) for number in citations) + "]")
        self._set_run_font(run, getattr(self, "_body_size_pt", 12), MUTED_HEX)

    @staticmethod
    def _mark_table_row_no_split(row) -> None:
        """禁止单个表格行跨页拆分。"""
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    @staticmethod
    def _mark_table_header(row) -> None:
        """标记重复表头，让跨页表格保留列名。"""
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            header = OxmlElement("w:tblHeader")
            header.set(qn("w:val"), "true")
            tr_pr.append(header)

    @staticmethod
    def _set_table_widths(table, widths: tuple[float, ...]) -> None:
        """为正式论文表格设置稳定列宽，避免 Word 自动挤压正文。"""

        for row in table.rows:
            for index, width in enumerate(widths):
                if index < len(row.cells):
                    row.cells[index].width = Inches(width)
