"""SPEC 0010 Word 模板支持测试。

覆盖：
- WordRenderer.render_with_template：模板渲染、占位符替换、章节循环、降级策略
- outlines API：上传/获取/删除/下载 Word 模板
- generate_word 端点返回 template_used 字段
"""

import io
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

from app.core.errors import AppError
from app.infrastructure.database.engine import Base
from app.infrastructure.renderers.word_renderer import WordRenderer
from app.main import app
from app.modules.outlines.models import Outline, WordTemplate
from app.modules.outlines.status import OutlineStatus
from app.modules.projects.models import Project
from app.modules.projects.status import ProjectStatus


TEST_DB = "sqlite:///:memory:"


SAMPLE_SECTIONS = [
    {
        "id": "sec_001",
        "title": "实验目的",
        "content": "分析胃病数据分布特征",
        "source_type": "REQUIREMENT",
        "source_ids": ["plan_001"],
    },
    {
        "id": "sec_002",
        "title": "实验结果",
        "content": "执行成功",
        "source_type": "EXECUTION",
        "source_ids": ["run_001"],
    },
]


# --- 渲染器测试 ---


def _create_template_docx(path: Path, with_sections_block: bool = True):
    """创建一个测试用 .docx 模板。"""
    doc = Document()
    doc.add_paragraph("课题：{{project_topic}}")
    doc.add_paragraph("项目：{{project_name}}")
    doc.add_paragraph("日期：{{generated_date}}")

    if with_sections_block:
        doc.add_paragraph("{{#sections}}")
        doc.add_paragraph("章节：{{section_title}}")
        doc.add_paragraph("内容：{{section_content}}")
        doc.add_paragraph("{{/sections}}")

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def test_render_with_template_replaces_cover_vars(tmp_path):
    """模板中的封面变量被正确替换。"""
    template_path = tmp_path / "template.docx"
    _create_template_docx(template_path, with_sections_block=False)

    output_path = tmp_path / "output.docx"
    renderer = WordRenderer()
    renderer.render_with_template(
        template_path=str(template_path),
        project_name="胃病数据分析",
        project_topic="胃病发病率研究",
        outline_sections=SAMPLE_SECTIONS,
        execution_artifacts=[],
        output_path=str(output_path),
    )

    doc = Document(str(output_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "胃病数据分析" in all_text
    assert "胃病发病率研究" in all_text
    assert "{{project_name}}" not in all_text
    assert "{{project_topic}}" not in all_text


def test_render_with_template_renders_sections_block(tmp_path):
    """模板中的章节循环块被正确渲染。"""
    template_path = tmp_path / "template.docx"
    _create_template_docx(template_path, with_sections_block=True)

    output_path = tmp_path / "output.docx"
    renderer = WordRenderer()
    renderer.render_with_template(
        template_path=str(template_path),
        project_name="测试项目",
        project_topic="测试课题",
        outline_sections=SAMPLE_SECTIONS,
        execution_artifacts=[],
        output_path=str(output_path),
    )

    doc = Document(str(output_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    # 两个章节都应出现
    assert "实验目的" in all_text
    assert "分析胃病数据分布特征" in all_text
    assert "实验结果" in all_text
    assert "执行成功" in all_text


def test_render_with_template_file_missing_raises_error(tmp_path):
    """模板文件不存在时抛出 WORD_TEMPLATE_FILE_MISSING。"""
    renderer = WordRenderer()
    with pytest.raises(AppError) as exc_info:
        renderer.render_with_template(
            template_path=str(tmp_path / "nonexistent.docx"),
            project_name="test",
            project_topic="test",
            outline_sections=[],
            execution_artifacts=[],
            output_path=str(tmp_path / "out.docx"),
        )
    assert exc_info.value.code == "WORD_TEMPLATE_FILE_MISSING"


def test_render_with_template_parse_failed_raises_error(tmp_path):
    """模板文件无法被 python-docx 打开时抛出 WORD_TEMPLATE_PARSE_FAILED。"""
    bad_path = tmp_path / "bad.docx"
    bad_path.write_bytes(b"not a docx file")

    renderer = WordRenderer()
    with pytest.raises(AppError) as exc_info:
        renderer.render_with_template(
            template_path=str(bad_path),
            project_name="test",
            project_topic="test",
            outline_sections=[],
            execution_artifacts=[],
            output_path=str(tmp_path / "out.docx"),
        )
    assert exc_info.value.code == "WORD_TEMPLATE_PARSE_FAILED"


def test_render_with_template_section_block_invalid(tmp_path):
    """模板含 {{#sections}} 但缺少 {{/sections}} 时抛出错误。"""
    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("{{#sections}}")
    doc.add_paragraph("{{section_title}}")
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(template_path))

    renderer = WordRenderer()
    with pytest.raises(AppError) as exc_info:
        renderer.render_with_template(
            template_path=str(template_path),
            project_name="test",
            project_topic="test",
            outline_sections=SAMPLE_SECTIONS,
            execution_artifacts=[],
            output_path=str(tmp_path / "out.docx"),
        )
    assert exc_info.value.code == "WORD_TEMPLATE_SECTION_BLOCK_INVALID"


def test_render_with_template_no_sections_block(tmp_path):
    """无循环块的模板只替换封面变量，并追加默认章节渲染。"""
    template_path = tmp_path / "template.docx"
    _create_template_docx(template_path, with_sections_block=False)

    output_path = tmp_path / "output.docx"
    renderer = WordRenderer()
    renderer.render_with_template(
        template_path=str(template_path),
        project_name="测试项目",
        project_topic="测试课题",
        outline_sections=SAMPLE_SECTIONS,
        execution_artifacts=[],
        output_path=str(output_path),
    )

    doc = Document(str(output_path))
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "测试项目" in all_text
    assert "测试课题" in all_text


# --- API 测试 ---


@pytest.fixture
def client(monkeypatch, tmp_path):
    """TestClient + 内存 SQLite + 受控工作区。"""
    monkeypatch.setenv("PROJECT_DATA_ROOT", str(tmp_path / "projects"))
    engine = create_engine(
        TEST_DB,
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    TestingSessionLocal = sessionmaker(bind=engine)

    from app.api.routers import projects as project_router
    from app.api.routers import outlines as outlines_router

    monkeypatch.setattr(project_router, "SessionLocal", TestingSessionLocal)
    monkeypatch.setattr(outlines_router, "SessionLocal", TestingSessionLocal)

    with TestClient(app) as test_client:
        yield test_client

    Base.metadata.drop_all(bind=engine)


def _create_project(client: TestClient) -> str:
    """创建项目并返回 project_id。"""
    response = client.post(
        "/api/projects",
        json={"name": "测试项目", "topic": "测试课题"},
    )
    assert response.status_code == 200
    return response.json()["id"]


def _make_docx_bytes() -> bytes:
    """生成一个有效的 .docx 文件字节流。"""
    doc = Document()
    doc.add_paragraph("模板内容 {{project_name}}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_upload_word_template_success(client):
    """上传 .docx 模板成功，返回模板信息。"""
    project_id = _create_project(client)
    docx_bytes = _make_docx_bytes()

    response = client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("template.docx", docx_bytes, "application/octet-stream")},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["project_id"] == project_id
    assert data["original_filename"] == "template.docx"
    assert data["file_size_bytes"] > 0
    assert len(data["content_hash"]) == 64  # SHA-256


def test_upload_word_template_non_docx_rejected(client):
    """上传非 .docx 文件返回 WORD_TEMPLATE_FILE_UNSUPPORTED。"""
    project_id = _create_project(client)

    response = client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("template.txt", b"hello", "text/plain")},
    )
    assert response.status_code == 400
    assert response.json()["error"]["code"] == "WORD_TEMPLATE_FILE_UNSUPPORTED"


def test_upload_word_template_too_large(client, monkeypatch):
    """上传超过大小限制的模板返回 WORD_TEMPLATE_TOO_LARGE。"""
    monkeypatch.setenv("WORD_TEMPLATE_MAX_SIZE_BYTES", "100")
    project_id = _create_project(client)
    large_bytes = b"x" * 200

    response = client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("big.docx", large_bytes, "application/octet-stream")},
    )
    assert response.status_code == 413
    assert response.json()["error"]["code"] == "WORD_TEMPLATE_TOO_LARGE"


def test_get_word_template_returns_template(client):
    """已上传模板时返回模板信息。"""
    project_id = _create_project(client)
    docx_bytes = _make_docx_bytes()

    client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("t.docx", docx_bytes, "application/octet-stream")},
    )
    response = client.get(f"/api/projects/{project_id}/word-template")
    assert response.status_code == 200
    data = response.json()
    assert data["original_filename"] == "t.docx"


def test_get_word_template_returns_null_when_no_template(client):
    """无模板时返回 null。"""
    project_id = _create_project(client)
    response = client.get(f"/api/projects/{project_id}/word-template")
    assert response.status_code == 200
    assert response.json() is None


def test_delete_word_template_success(client):
    """删除已上传模板返回 204。"""
    project_id = _create_project(client)
    docx_bytes = _make_docx_bytes()

    client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("t.docx", docx_bytes, "application/octet-stream")},
    )
    response = client.delete(f"/api/projects/{project_id}/word-template")
    assert response.status_code == 204

    # 确认已删除
    response = client.get(f"/api/projects/{project_id}/word-template")
    assert response.json() is None


def test_delete_word_template_not_found(client):
    """删除不存在的模板返回 404。"""
    project_id = _create_project(client)
    response = client.delete(f"/api/projects/{project_id}/word-template")
    assert response.status_code == 404
    assert response.json()["error"]["code"] == "WORD_TEMPLATE_NOT_FOUND"


def test_upload_word_template_replaces_existing(client):
    """重新上传模板覆盖旧模板。"""
    project_id = _create_project(client)
    docx_bytes = _make_docx_bytes()

    # 第一次上传
    r1 = client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("v1.docx", docx_bytes, "application/octet-stream")},
    )
    assert r1.status_code == 200

    # 第二次上传（替换）
    r2 = client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("v2.docx", docx_bytes, "application/octet-stream")},
    )
    assert r2.status_code == 200
    assert r2.json()["original_filename"] == "v2.docx"

    # 确认只有一个模板
    r3 = client.get(f"/api/projects/{project_id}/word-template")
    assert r3.json()["original_filename"] == "v2.docx"


def test_download_word_template_success(client):
    """下载已上传的模板文件。"""
    project_id = _create_project(client)
    docx_bytes = _make_docx_bytes()

    client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("t.docx", docx_bytes, "application/octet-stream")},
    )
    response = client.get(f"/api/projects/{project_id}/word-template/download")
    assert response.status_code == 200
    assert len(response.content) > 0


def test_download_word_template_not_found(client):
    """下载不存在的模板返回 404。"""
    project_id = _create_project(client)
    response = client.get(f"/api/projects/{project_id}/word-template/download")
    assert response.status_code == 404
    assert response.json()["error"]["code"] == "WORD_TEMPLATE_NOT_FOUND"


def test_generate_word_returns_template_used_false(client, monkeypatch, tmp_path):
    """无模板时 generate_word 的响应 template_used 为 false。"""
    monkeypatch.setenv("PROJECT_DATA_ROOT", str(tmp_path / "projects"))
    project_id = _create_project(client)

    # 创建已确认大纲
    from app.api.routers import outlines as outlines_router
    from sqlalchemy.orm import sessionmaker

    engine = create_engine(
        TEST_DB,
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    # 使用已注入的 SessionLocal
    db = outlines_router.SessionLocal()
    try:
        project = db.query(Project).filter(Project.id == project_id).first()
        project.status = ProjectStatus.OUTLINE_CONFIRMED.value

        import json
        outline = Outline(
            id="ol_test1",
            project_id=project_id,
            sections_json=json.dumps(SAMPLE_SECTIONS),
            status=OutlineStatus.CONFIRMED.value,
            candidate_source="LOCAL_RULE",
            code_version=1,
        )
        db.add(outline)
        db.commit()
        outline_id = outline.id
    finally:
        db.close()

    response = client.post(
        f"/api/projects/{project_id}/outline/{outline_id}/word/generate"
    )
    assert response.status_code == 201
    data = response.json()
    assert "template_used" in data
    assert data["template_used"] is False


def test_generate_word_returns_template_used_true(client, monkeypatch, tmp_path):
    """有模板时 generate_word 的响应 template_used 为 true。"""
    monkeypatch.setenv("PROJECT_DATA_ROOT", str(tmp_path / "projects"))
    project_id = _create_project(client)

    # 上传模板
    docx_bytes = _make_docx_bytes()
    client.post(
        f"/api/projects/{project_id}/word-template",
        files={"file": ("t.docx", docx_bytes, "application/octet-stream")},
    )

    # 创建已确认大纲
    from app.api.routers import outlines as outlines_router

    db = outlines_router.SessionLocal()
    try:
        project = db.query(Project).filter(Project.id == project_id).first()
        project.status = ProjectStatus.OUTLINE_CONFIRMED.value

        import json
        outline = Outline(
            id="ol_test2",
            project_id=project_id,
            sections_json=json.dumps(SAMPLE_SECTIONS),
            status=OutlineStatus.CONFIRMED.value,
            candidate_source="LOCAL_RULE",
            code_version=1,
        )
        db.add(outline)
        db.commit()
        outline_id = outline.id
    finally:
        db.close()

    response = client.post(
        f"/api/projects/{project_id}/outline/{outline_id}/word/generate"
    )
    assert response.status_code == 201
    data = response.json()
    assert "template_used" in data
    assert data["template_used"] is True
