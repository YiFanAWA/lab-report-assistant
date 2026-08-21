"""SPEC 0035/0036 大样本论文解读案例的口径、来源与深度分析测试。"""

import json
from pathlib import Path


def test_spec0035_fixture_has_large_dataset_and_paper_sources():
    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    data_path = base / "data" / "diabetic_data.csv"
    paper_path = base / "sources" / "strack_2014_hba1c_readmission.pdf"
    xml_path = base / "sources" / "strack_2014_hba1c_readmission.xml"
    manifest_path = base / "sources" / "source_manifest.json"

    assert data_path.exists()
    assert paper_path.exists()
    assert xml_path.exists()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    assert manifest["paper"]["doi"] == "10.1155/2014/781670"
    assert manifest["dataset"]["license"] == "CC BY 4.0"


def test_spec0035_analysis_summary_preserves_paper_vs_reanalysis_metrics():
    summary_path = (
        Path(__file__).parents[1]
        / "dev-docs"
        / "e2e-screenshots"
        / "spec0035_paper_review"
        / "analysis_summary.json"
    )
    summary = json.loads(summary_path.read_text(encoding="utf-8"))

    assert summary["total_records"] == 101766
    assert summary["unique_patients"] == 71518
    assert summary["repeated_patient_records"] == 30248
    assert summary["index_records"] == 71518
    assert summary["death_hospice_index_records"] == 1545
    assert summary["primary_analysis_records"] == 69973
    assert summary["columns"] == 50
    assert summary["paper_final_sample"] == 69984
    assert summary["paper_reported_a1c_measurement_rate"] == 0.184
    assert summary["early_rate_measured"] < summary["early_rate_unmeasured"]
    assert summary["a1c_status_counts"] == {"measured": 12845, "not_measured": 57128, "missing_or_unknown": 0}
    assert summary["model_contract"]["analysis_position"] == "教学性论文复核报告，不是独立研究论文"
    assert summary["model_contract"]["standard_errors"].startswith("按 patient_nbr 聚类")
    assert summary["interaction_joint_p_value"] > 0


def test_spec0036_analysis_summary_contains_depth_outputs():
    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    summary = json.loads((base / "analysis_summary.json").read_text(encoding="utf-8"))
    charts = base / "charts"

    assert summary["artifact_count"] >= 11
    assert summary["uci_features"] == 47
    assert summary["risk_difference"] < 0
    assert len(summary["risk_difference_ci"]) == 2
    assert summary["top_missingness"]["weight"] > 0.9
    assert len(summary["logistic_reanalysis"]) >= 20
    assert {"a1c_measured", "a1c_x_Diabetes", "a1c_x_Respiratory"} <= {row["term"] for row in summary["logistic_reanalysis"]}
    assert set(summary["sensitivity_analysis"]) == {"clustered_repeated_records", "including_death_hospice"}
    assert {
        plan["chart_kind"] for plan in summary["chart_plans"].values()
    } >= {"flow", "stacked_composition", "dumbbell", "point_ci", "ordered_line", "forest"}
    for filename in (
        "sample_cohort_flow.png",
        "missingness_top8.png",
        "primary_effect_summary.csv",
        "logistic_forest.png",
        "logistic_reanalysis.csv",
        "logistic_reanalysis_full.csv",
        "a1c_status_summary.csv",
        "primary_diagnosis_summary.csv",
        "variable_coding.csv",
        "sensitivity_analysis.csv",
        "primary_diagnosis_a1c_interaction.png",
    ):
        assert (charts / filename).exists()


def test_spec0037_chart_plans_are_traceable_to_real_artifacts():
    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    summary = json.loads((base / "analysis_summary.json").read_text(encoding="utf-8"))
    chart_plans = summary["chart_plans"]
    assert summary["execution_run_id"] == "spec0040_argumentation"
    for name, plan in chart_plans.items():
        assert (base / "charts" / plan["file_name"]).exists(), name
        assert plan["encoding"]
        assert plan["rationale"]


def test_spec0039_semantic_figures_are_traceable_and_non_hierarchical():
    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    summary = json.loads((base / "analysis_summary.json").read_text(encoding="utf-8"))
    figure_plans = summary["figure_plans"]
    expected = {
        "研究证据链图": ("evidence_chain", "evidence_chain.png"),
        "数据处理管线图": ("data_pipeline", "data_pipeline.png"),
        "变量关系图": ("relationship_graph", "variable_relationship.png"),
    }
    assert set(expected) <= set(figure_plans)
    for name, (kind, filename) in expected.items():
        plan = figure_plans[name]
        assert plan["figure_kind"] == kind
        assert plan["source_ids"]
        assert plan["execution_run_ids"] == ["spec0040_argumentation"]
        assert (base / "charts" / filename).exists()
    relationship = figure_plans["变量关系图"]
    assert any(edge["relation"] == "associational" for edge in relationship["edges"])
    assert "不代表因果" in relationship["note"]
    for plan in figure_plans.values():
        assert plan["argument"]["claim"]
        assert plan["argument"]["evidence_refs"]
        assert plan["argument"]["result"]
        assert plan["argument"]["boundary"]


def test_spec0041_heterogeneous_portfolio_contains_matrix_and_rejected_candidates():
    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    summary = json.loads((base / "analysis_summary.json").read_text(encoding="utf-8"))
    portfolio = summary["figure_portfolio_plan"]
    families = {figure["visual_family"] for figure in portfolio["figures"]}
    assert {"evidence_argument", "process", "relationship", "matrix", "statistical"} <= families
    assert (base / "charts" / "paper_local_comparison_matrix.png").exists()
    assert any(candidate["name"] == "质量热力图" for candidate in portfolio["rejected_candidates"])


def test_spec0040_evidence_chain_is_a_four_panel_argument_figure():
    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    summary = json.loads((base / "analysis_summary.json").read_text(encoding="utf-8"))
    plan = summary["figure_plans"]["研究证据链图"]
    assert plan["figure_kind"] == "evidence_chain"
    assert len(plan["nodes"]) >= 8
    assert len(plan["panel_labels"]) == 4
    assert {edge["relation"] for edge in plan["edges"]} >= {
        "supports", "contains", "produces", "compared_with", "bounded_by"
    }
    assert plan["argument"]["evidence_refs"]
    assert plan["argument"]["body_reference"] == "见第 2.1 节"
    assert plan["source_ids"] == ["paper:PMC3996476", "dataset:UCI-296"]
    assert plan["execution_run_ids"] == ["spec0040_argumentation"]
    figure = base / "charts" / "evidence_chain.png"
    assert figure.exists()
    # 多面板期刊图应明显宽于单行流程图，且保持高 DPI 输出。
    from PIL import Image

    with Image.open(figure) as image:
        assert image.width >= 3000
        assert image.height >= 1800


def test_spec0036_deliverables_are_expanded_and_reopenable():
    from docx import Document
    from pptx import Presentation

    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    document = Document(str(base / "spec0035_paper_review.docx"))
    assert len(document.inline_shapes) >= 9
    document_text = "\n".join(p.text for p in document.paragraphs)
    assert "教学性多变量模型与交互分析" in document_text
    assert "主要诊断分层与 HbA1c 交互" in document_text
    assert "非独立研究论文" in document_text
    # SPEC 0039 新增证据链、数据管线和变量关系三页，故答辩稿由 13 页扩展为 16 页。
    for filename in ("spec0035_paper_review.pptx", "spec0035_sjtu_paper_review.pptx"):
        deck = Presentation(str(base / filename))
        assert 16 <= len(deck.slides) <= 18
        assert round(deck.slide_width / 914400, 2) == 13.33
        assert round(deck.slide_height / 914400, 2) == 7.5


def test_spec0043_publication_outputs_are_named_and_bound_to_one_docx_pdf_chain():
    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    manifest = json.loads((base / "publication_manifest.json").read_text(encoding="utf-8"))
    assert manifest["spec"] == "0043"
    assert manifest["deliverables"]["docx"] == "spec0043_publication.docx"
    assert manifest["deliverables"]["pdf"] == "spec0043_publication.pdf"
    assert manifest["deliverables"]["pptx"] == "spec0043_publication.pptx"
    assert manifest["pdf_binding"]["exporter"] == "DocxPdfExporter"
    assert manifest["pdf_binding"]["source_docx"] == "spec0043_publication.docx"
    assert (base / manifest["deliverables"]["docx"]).exists()
    assert (base / manifest["deliverables"]["pdf"]).read_bytes().startswith(b"%PDF-")
    assert (base / manifest["deliverables"]["pptx"]).exists()
    model_csv = base / "charts" / "logistic_reanalysis.csv"
    assert "P=0" not in model_csv.read_text(encoding="utf-8")
    from pypdf import PdfReader

    pdf_text = "\n".join((page.extract_text() or "") for page in PdfReader(str(base / manifest["deliverables"]["pdf"])).pages)
    assert "STROBE" in pdf_text
    assert "Cameron" in pdf_text
    assert "Ibrahim" in pdf_text
    assert "简化多变量模型复核" not in pdf_text


def test_spec0043_body_projection_excludes_engineering_inventory():
    from docx import Document

    base = Path(__file__).parents[1] / "dev-docs" / "e2e-screenshots" / "spec0035_paper_review"
    text = "\n".join(p.text for p in Document(str(base / "spec0043_publication.docx")).paragraphs)
    body = text.split("附录：执行产物索引", 1)[0]
    forbidden = ("SHA-256", "sha256", "JSON path", "JSON 路径", "执行批次：", "file_path")
    assert not any(token in body for token in forbidden)
    assert "研究问题与论文假设" in body
    assert "结论与可复核证据链" in body