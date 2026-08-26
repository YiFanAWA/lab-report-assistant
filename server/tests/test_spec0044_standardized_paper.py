"""SPEC 0044 reader-first publication gate."""

from __future__ import annotations

import zipfile

from docx import Document
from docx.shared import RGBColor
from PIL import Image

from app.infrastructure.renderers.word_renderer import WordRenderer
from app.modules.outlines.document_planner import plan_thesis_document


def _sections() -> list[dict]:
    return [
        {
            "title": "Research question",
            "content": "This study compares the primary outcome between two public-data groups.",
            "source_type": "REQUIREMENT",
            "manuscript_role": "introduction",
            "research_question": "Do the two groups differ descriptively?",
            "source_ids": ["paper"],
            "subsections": [
                {
                    "title": "Operational definition",
                    "content": "The primary outcome is defined before comparison.",
                },
            ],
        },
        {
            "title": "Study population and data source",
            "content": "The study population comes from a public dataset.",
            "source_type": "DATASET",
            "manuscript_role": "methods",
            "source_ids": ["dataset"],
        },
        {
            "title": "Statistical methods",
            "content": "Descriptive statistics, interval estimation, and stratified comparisons were used.",
            "source_type": "ANALYSIS",
            "manuscript_role": "methods",
            "source_ids": ["paper"],
        },
        {
            "title": "Primary result",
            "content": "The primary result is shown in the figure and table.",
            "source_type": "EXECUTION",
            "manuscript_role": "results",
            "artifact_group": "primary",
            "source_ids": ["run"],
            "figure_takeaway": "The groups differ descriptively, but this does not establish causality.",
        },
        {
            "title": "Discussion and limitations",
            "content": "Observational data contain confounding and cohort-definition limits.",
            "source_type": "SUMMARY",
            "manuscript_role": "discussion",
            "source_ids": ["paper"],
        },
    ]


def test_reader_first_formal_render_hides_engineering_projection(tmp_path):
    image = tmp_path / "result.png"
    Image.new("RGB", (640, 360), "white").save(image)
    table = tmp_path / "summary.csv"
    table.write_text("group,n,rate\nmeasured,100,10%\nunmeasured,100,15%\n", encoding="utf-8")
    output = tmp_path / "paper.docx"

    WordRenderer().render(
        project_name="Paper project",
        project_topic="Public data analysis",
        outline_sections=_sections(),
        execution_artifacts=[
            {
                "name": "result.png",
                "artifact_type": "CHART_PNG",
                "file_path": str(image),
                "execution_run_id": "run-0044",
                "artifact_group": "primary",
                "figure_caption": "Primary outcome comparison",
                "figure_layout_profile": "side_by_side",
                "figure_note": "Note: intervals express estimation uncertainty.",
                "figure_body_reference": "primary result",
                "scientific_asset_ids": ["asset-1"],
                "scientific_asset_image_sha256": "deadbeef",
                "figure_argument": {
                    "claim": "engineering claim",
                    "method": "engineering method",
                    "result": "engineering result",
                    "boundary": "engineering boundary",
                },
            },
            {
                "name": "secondary.png",
                "artifact_type": "CHART_PNG",
                "file_path": str(image),
                "execution_run_id": "run-0044",
                "artifact_group": "primary",
                "figure_caption": "Secondary outcome comparison",
                "figure_layout_profile": "side_by_side",
                "figure_note": "Note: the secondary panel is descriptive.",
            },
            {
                "name": "summary.csv",
                "artifact_type": "TABLE_CSV",
                "file_path": str(table),
                "execution_run_id": "run-0044",
                "artifact_group": "primary",
                "table_caption": "Primary outcome summary",
            },
        ],
        output_path=str(output),
        config={
            "document_profile": "formal_academic",
            "formal_metadata": {
                "作者": "作者甲",
                "单位": "单位乙",
                "study_type": "public data descriptive analysis",
                "execution_run_id": "run-0044",
            },
            "abstract_sections": {
                "目的": "明确研究问题。",
                "方法": "采用描述性分析。",
                "结果": "两组存在差异。",
                "结论": "结果仅支持方向性判断。",
            },
            "abstract_sections_en": {
                "Purpose": "Define the research question.",
                "Methods": "Use descriptive analysis.",
                "Results": "The groups differ descriptively.",
                "Conclusion": "Interpretation remains limited.",
            },
            "reference_catalog": {
                "paper": "Example paper.",
                "dataset": "Example dataset.",
            },
        },
    )

    document = Document(output)
    paragraphs = document.paragraphs
    body_text = "\n".join(paragraph.text for paragraph in paragraphs)
    forbidden = (
        "Artifact Source",
        "Figure Lead",
        "\u9644\u5f55\uff1a\u6267\u884c\u4ea7\u7269\u7d22\u5f15",
        "\u6267\u884c\u6279\u6b21\uff1a",
        "\u6267\u884c\u6279\u6b21",
        "\u79d1\u7814\u8d44\u4ea7",
        "SHA-256",
        "\u6e32\u67d3\u8ffd\u6eaf",
        "\u8bba\u8bc1\uff1a",
        "\u6b63\u6587\u5f15\u7528\uff1a",
        "\u6765\u6e90\uff1a\u6267\u884c\u4ea7\u7269",
    )
    assert not any(token in body_text for token in forbidden)
    assert not any(
        paragraph.style.name in {"Artifact Source", "Figure Lead"}
        for paragraph in paragraphs
    )
    assert any("Secondary outcome comparison" in paragraph.text for paragraph in paragraphs)
    assert any("图 1  Primary outcome comparison" in paragraph.text for paragraph in paragraphs)
    assert any("图 2  Secondary outcome comparison" in paragraph.text for paragraph in paragraphs)
    assert any("表 1  Primary outcome summary" in paragraph.text for paragraph in paragraphs)
    assert not any("图 1-1" in paragraph.text or "图 2-1" in paragraph.text for paragraph in paragraphs)
    assert not any("相关结果见图" in paragraph.text for paragraph in paragraphs)
    assert any("作者：作者甲" in paragraph.text for paragraph in paragraphs)
    assert any("单位：单位乙" in paragraph.text for paragraph in paragraphs)
    assert any("Primary outcome comparison" in paragraph.text for paragraph in paragraphs)
    assert any("Primary outcome summary" in paragraph.text for paragraph in paragraphs)
    assert any("摘要" == paragraph.text.strip() for paragraph in paragraphs)
    assert any("ABSTRACT" == paragraph.text.strip() for paragraph in paragraphs)
    assert any("目的：明确研究问题。" in paragraph.text for paragraph in paragraphs)
    assert any("Purpose: Define the research question." in paragraph.text for paragraph in paragraphs)
    assert any("图目录" == paragraph.text.strip() for paragraph in paragraphs)
    assert any("表目录" == paragraph.text.strip() for paragraph in paragraphs)
    assert not any("图表目录" in paragraph.text for paragraph in paragraphs)
    assert any("图 1  Primary outcome comparison" in paragraph.text for paragraph in paragraphs)
    assert any("表 1  Primary outcome summary" in paragraph.text for paragraph in paragraphs)
    assert not any(paragraph.text.strip() == "????" for paragraph in paragraphs)
    assert document.styles["Body Text"].font.size.pt == 10.5
    assert document.styles["Heading 1"].font.size.pt == 16.0
    assert document.styles["Caption"].font.size.pt == 9.0
    assert all(
        run.font.size is not None and run.font.size.pt == 9.0
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    )
    toc_entry = next(
        paragraph for paragraph in paragraphs if paragraph.text.startswith("1.1 ")
    )
    assert all(
        run.font.size is not None and run.font.size.pt == 10.5
        for run in toc_entry.runs
        if run.text.strip()
    )
    assert document.styles["Heading 1"].font.color.rgb == RGBColor(0x22, 0x22, 0x22)
    assert len(document.sections) >= 3

    with zipfile.ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "PAGEREF chapter_1" in xml
    assert "PAGEREF section_1_1" in xml
    assert "PAGEREF section_1_1_1" in xml
    assert "SEQ Figure" in xml
    assert "SEQ Table" in xml
    assert "REF fig_1" in xml
    assert "PAGEREF fig_1" in xml
    assert "REF fig_2" in xml
    assert "PAGEREF fig_2" in xml
    assert "PAGEREF tbl_1" in xml
    assert "w:tblHeader" in xml
    assert "w:cantSplit" in xml


def test_manuscript_plan_reader_projection_removes_trace_wording_without_mutating_input():
    raw_sections = _sections()
    batch = "".join(chr(v) for v in (0x6267, 0x884c, 0x6279, 0x6b21))
    artifact = "".join(chr(v) for v in (0x6267, 0x884c, 0x4ea7, 0x7269))
    raw_sections[3]["content"] += " " + batch
    raw_sections[3]["figure_takeaway"] = artifact + " and " + batch + " are internal wording."
    plan = plan_thesis_document(
        "Public data analysis",
        raw_sections,
        [{"artifact_type": "CHART_PNG"}],
        formal=True,
        reference_catalog={"paper": "Example paper.", "dataset": "Example dataset."},
    )
    projected = repr(plan.chapters)
    assert batch in raw_sections[3]["content"]
    assert artifact in raw_sections[3]["figure_takeaway"]
    assert batch not in projected
    assert artifact not in projected
